# -*- coding: utf-8 -*-

# =================================================================================
# osintgpt
#
# Author: @estebanpdl
#
# File: office.py
# Description: Word documents as markdown. The structure a writer applied —
#   headings, tables, lists — is what chunking later reads, so it is preserved
#   rather than flattened into a wall of text.
# =================================================================================

# import submodules
from pathlib import Path

# type hints
from typing import List, Union

# Word's own heading styles, mapped to the markdown the chunker understands.
# A document written without them yields plain paragraphs, which is the
# unstructured case and handled the same as any other.
_HEADING_STYLES = ('Heading', 'Title', 'Subtitle')


# read a Word document as markdown
def extract_docx(path: Union[str, Path]) -> str:
    '''
    Read a .docx into markdown.

    Headings become `#` levels and tables become pipe rows, because those are
    the two structures chunking acts on: a heading opens a section, and a
    table is kept whole with its header.

    Args:
        path (Union[str, Path]): Document to read.

    Raises:
        ImportError: If python-docx is not installed.

    Returns:
        str: The document as markdown.
    '''
    try:
        import docx
    except ImportError as error:
        raise ImportError(
            "reading Word documents needs the 'python-docx' package, which "
            'osintgpt requires: reinstall with pip install '
            '--force-reinstall osintgpt'
        ) from error

    document = docx.Document(str(path))
    blocks: List[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        level = _heading_level(paragraph)
        blocks.append(f'{"#" * level} {text}' if level else text)

    for table in document.tables:
        rendered = _render_table(table)
        if rendered:
            blocks.append(rendered)

    return '\n\n'.join(blocks)


def _heading_level(paragraph) -> int:
    '''
    The markdown level a paragraph's style implies, or 0 for body text.

    Title and Subtitle become the top levels; `Heading 3` becomes `###`. A
    document that styles nothing has no headings, which is ordinary.
    '''
    name = getattr(paragraph.style, 'name', '') or ''
    if not name.startswith(_HEADING_STYLES):
        return 0

    if name.startswith('Title'):
        return 1
    if name.startswith('Subtitle'):
        return 2

    digits = ''.join(c for c in name if c.isdigit())

    return min(int(digits), 6) if digits else 1


def _render_table(table) -> str:
    '''
    A Word table as markdown rows, with the rule that marks its first row as a
    header — which is what lets chunking repeat it when the table is split.
    '''
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        if any(cells):
            rows.append('| ' + ' | '.join(cells) + ' |')

    if not rows:
        return ''

    columns = len(table.columns)
    rule = '|' + '|'.join(['---'] * columns) + '|'

    return '\n'.join([rows[0], rule] + rows[1:])
